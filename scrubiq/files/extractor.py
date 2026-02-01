"""
Text extractors for various file formats.

Each extractor implements a common interface for extracting text content
from uploaded files, with OCR fallback for scanned documents.
"""

import csv
import io
import logging
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional, Dict, Any, TYPE_CHECKING

from ..constants import (
    MIN_NATIVE_TEXT_LENGTH,
    MAX_DOCUMENT_PAGES,
    MAX_SPREADSHEET_ROWS,
    MAX_DECOMPRESSED_SIZE,
    MAX_EXTRACTION_RATIO,
)

if TYPE_CHECKING:
    from .ocr import OCREngine, OCRResult
    from .temp_storage import SecureTempDir
    from .document_templates import DocumentType

logger = logging.getLogger(__name__)


@dataclass
class PageInfo:
    """Information about a single extracted page."""
    page_num: int
    text: str
    is_scanned: bool  # True if this page needed OCR
    ocr_result: Optional["OCRResult"] = None  # OCR data with coordinates
    temp_image_path: Optional[str] = None  # Path to rendered image in temp dir


@dataclass
class ExtractionResult:
    """Result of text extraction from a file."""
    text: str
    pages: int = 1
    needs_ocr: bool = False  # True if any page needed OCR
    ocr_pages: List[int] = field(default_factory=list)  # Which pages needed OCR
    warnings: List[str] = field(default_factory=list)  # Non-fatal issues
    confidence: float = 1.0  # Average OCR confidence (1.0 if no OCR)
    ocr_results: List["OCRResult"] = field(default_factory=list)  # OCR data with coordinates per page
    
    # Per-page tracking for multi-page visual redaction
    page_infos: List[PageInfo] = field(default_factory=list)
    
    # Temp directory containing rendered page images (for cleanup)
    temp_dir_path: Optional[str] = None
    
    # Document intelligence (from EnhancedOCRProcessor)
    document_type: Optional[str] = None  # e.g., "DRIVERS_LICENSE", "INSURANCE_MEDICARE"
    is_id_document: bool = False  # True if ID card, insurance card, etc.
    phi_fields: Optional[Dict[str, Any]] = None  # Pre-extracted PHI with categories
    enhanced_text: Optional[str] = None  # Layout-improved text (if different from raw)
    enhancements_applied: List[str] = field(default_factory=list)  # What processing was done
    
    @property
    def has_scanned_pages(self) -> bool:
        """Check if any pages are scanned (need visual redaction)."""
        return any(p.is_scanned for p in self.page_infos)
    
    @property
    def scanned_page_count(self) -> int:
        """Count of scanned pages."""
        return sum(1 for p in self.page_infos if p.is_scanned)
    
    @property
    def best_text(self) -> str:
        """Return enhanced text if available, otherwise raw text."""
        return self.enhanced_text if self.enhanced_text else self.text


class BaseExtractor(ABC):
    """Base class for format-specific extractors."""
    
    @abstractmethod
    def can_handle(self, content_type: str, extension: str) -> bool:
        """
        Check if this extractor handles the file type.
        
        Args:
            content_type: MIME type
            extension: File extension (lowercase, with dot)
            
        Returns:
            True if this extractor can process the file
        """
        pass
    
    @abstractmethod
    def extract(self, content: bytes, filename: str) -> ExtractionResult:
        """
        Extract text from file content.
        
        Args:
            content: Raw file bytes
            filename: Original filename (for logging/extension detection)
            
        Returns:
            ExtractionResult with extracted text and metadata
        """
        pass


class PDFExtractor(BaseExtractor):
    """
    PDF text extractor using PyMuPDF with document intelligence.
    
    Handles both text-layer PDFs and scanned documents:
    - Extracts text layer if available (native PDF)
    - Falls back to OCR for pages without text (scanned PDF)
    - Uses EnhancedOCRProcessor for document type detection and PHI extraction
    
    For visual redaction support:
    - Renders scanned pages to temp files at 150 DPI
    - Uses extract_with_coordinates for OCR bounding boxes
    - Tracks per-page scanned status in PageInfo
    """
    
    # DPI for rendering scanned pages (150 = good balance of quality and size)
    RENDER_DPI = 150
    
    def __init__(
        self, 
        ocr_engine: Optional["OCREngine"] = None,
        temp_dir: Optional["SecureTempDir"] = None,
        enable_enhanced_processing: bool = True,
    ):
        """
        Initialize PDF extractor.
        
        Args:
            ocr_engine: OCR engine for scanned pages
            temp_dir: Secure temp directory for page images (optional)
            enable_enhanced_processing: Use EnhancedOCRProcessor for document intelligence
        """
        self.ocr_engine = ocr_engine
        self.temp_dir = temp_dir
        self.enable_enhanced_processing = enable_enhanced_processing
        self._enhanced_processor = None
    
    @property
    def enhanced_processor(self):
        """Lazy-load EnhancedOCRProcessor."""
        if self._enhanced_processor is None and self.enable_enhanced_processing:
            try:
                from .enhanced_ocr import EnhancedOCRProcessor
                self._enhanced_processor = EnhancedOCRProcessor()
            except Exception as e:
                logger.warning(f"Could not initialize EnhancedOCRProcessor: {e}")
                self.enable_enhanced_processing = False
        return self._enhanced_processor
    
    def can_handle(self, content_type: str, extension: str) -> bool:
        return (
            content_type == "application/pdf" or 
            extension == ".pdf"
        )
    
    def extract(
        self, 
        content: bytes, 
        filename: str,
        save_scanned_pages: bool = True,
    ) -> ExtractionResult:
        """
        Extract text from PDF with document intelligence for scanned pages.
        
        Args:
            content: PDF bytes
            filename: Original filename
            save_scanned_pages: If True, save rendered images of scanned pages
                               to temp directory for visual redaction
                               
        Returns:
            ExtractionResult with per-page info and document intelligence
        """
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise ImportError("PyMuPDF not installed. Run: pip install pymupdf")
        
        doc = fitz.open(stream=content, filetype="pdf")
        
        pages_text = []
        page_infos = []
        ocr_pages = []
        ocr_results = []
        ocr_confidences = []
        warnings = []
        
        # Document intelligence (from first scanned page)
        document_type = None
        is_id_document = False
        phi_fields = None
        enhanced_texts = []
        enhancements = []
        first_scanned_processed = False
        
        try:
            for i, page in enumerate(doc):
                # Early exit if page limit exceeded (prevents DoS via large PDFs)
                if i >= MAX_DOCUMENT_PAGES:
                    logger.warning(f"PDF exceeds {MAX_DOCUMENT_PAGES} page limit, truncating")
                    warnings.append(f"Document truncated at {MAX_DOCUMENT_PAGES} pages")
                    break

                # Try to extract text layer
                native_text = page.get_text().strip()
                
                # Check if this page has meaningful native text
                has_native_text = len(native_text) >= MIN_NATIVE_TEXT_LENGTH
                
                if has_native_text:
                    # Native text page - no visual redaction needed
                    pages_text.append(native_text)
                    enhanced_texts.append(native_text)
                    page_infos.append(PageInfo(
                        page_num=i,
                        text=native_text,
                        is_scanned=False,
                        ocr_result=None,
                        temp_image_path=None,
                    ))
                    logger.debug(f"Page {i+1}: native text ({len(native_text)} chars)")
                    
                elif self.ocr_engine and self.ocr_engine.is_available:
                    # Scanned page - needs OCR and visual redaction
                    logger.debug(f"Page {i+1}: scanned, using OCR")
                    
                    try:
                        # Render page to image
                        pix = page.get_pixmap(dpi=self.RENDER_DPI)
                        
                        # Convert to PIL Image and numpy array
                        from PIL import Image
                        import numpy as np
                        
                        img = Image.frombytes(
                            "RGB", 
                            [pix.width, pix.height], 
                            pix.samples
                        )
                        img_array = np.array(img)
                        
                        # Save to temp file if requested
                        temp_path = None
                        if save_scanned_pages and self.temp_dir:
                            # Convert to PNG bytes
                            img_buffer = io.BytesIO()
                            img.save(img_buffer, format='PNG')
                            img_bytes = img_buffer.getvalue()
                            
                            # Write to temp dir
                            temp_path_obj = self.temp_dir.write_page(i, img_bytes)
                            temp_path = str(temp_path_obj)
                        
                        # Get raw OCR with coordinates
                        ocr_result = self.ocr_engine.extract_with_coordinates(img)
                        
                        # Apply enhanced processing
                        enhanced_text = ocr_result.full_text
                        
                        if self.enable_enhanced_processing and self.enhanced_processor and ocr_result.blocks:
                            try:
                                enhanced_result = self.enhanced_processor.process(
                                    image=img_array,
                                    ocr_result=ocr_result,
                                    apply_document_cleaning=True,
                                )
                                
                                enhanced_text = enhanced_result.enhanced_text
                                
                                # Capture document type from first scanned page
                                if not first_scanned_processed:
                                    document_type = enhanced_result.document_type.name
                                    is_id_document = enhanced_result.is_id_card
                                    phi_fields = enhanced_result.phi_fields
                                    enhancements = enhanced_result.enhancements_applied
                                    first_scanned_processed = True
                                    
                                    logger.info(
                                        f"PDF document intelligence (page {i+1}): "
                                        f"type={document_type}, is_id={is_id_document}, "
                                        f"phi_fields={len(phi_fields) if phi_fields else 0}"
                                    )
                                
                            except Exception as e:
                                logger.warning(f"Enhanced processing failed for page {i+1}: {e}")
                        
                        pages_text.append(ocr_result.full_text)
                        enhanced_texts.append(enhanced_text)
                        ocr_pages.append(i)
                        ocr_results.append(ocr_result)
                        ocr_confidences.append(ocr_result.confidence)
                        
                        page_infos.append(PageInfo(
                            page_num=i,
                            text=enhanced_text,
                            is_scanned=True,
                            ocr_result=ocr_result,
                            temp_image_path=temp_path,
                        ))
                        
                    except Exception as e:
                        logger.warning(f"OCR failed for page {i+1}: {e}")
                        pages_text.append("")
                        enhanced_texts.append("")
                        warnings.append(f"OCR failed for page {i+1}: {e}")
                        page_infos.append(PageInfo(
                            page_num=i,
                            text="",
                            is_scanned=True,
                            ocr_result=None,
                            temp_image_path=None,
                        ))
                else:
                    # No OCR available
                    pages_text.append("")
                    enhanced_texts.append("")
                    page_infos.append(PageInfo(
                        page_num=i,
                        text="",
                        is_scanned=True,  # Assume scanned if no text
                        ocr_result=None,
                        temp_image_path=None,
                    ))
                    if not self.ocr_engine:
                        warnings.append(f"Page {i+1} is scanned but OCR not available")
            
            # Calculate average OCR confidence
            avg_confidence = 1.0
            if ocr_confidences:
                avg_confidence = sum(ocr_confidences) / len(ocr_confidences)
            
            return ExtractionResult(
                text="\n\n".join(pages_text),
                pages=len(doc),
                needs_ocr=len(ocr_pages) > 0,
                ocr_pages=ocr_pages,
                warnings=warnings,
                confidence=avg_confidence,
                ocr_results=ocr_results,
                page_infos=page_infos,
                temp_dir_path=str(self.temp_dir.path) if self.temp_dir and self.temp_dir.path else None,
                # Document intelligence
                document_type=document_type,
                is_id_document=is_id_document,
                phi_fields=phi_fields,
                enhanced_text="\n\n".join(enhanced_texts) if enhanced_texts else None,
                enhancements_applied=enhancements,
            )
            
        finally:
            doc.close()


class DOCXExtractor(BaseExtractor):
    """Word document extractor using python-docx."""
    
    def can_handle(self, content_type: str, extension: str) -> bool:
        return (
            content_type in (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "application/msword",
            ) or 
            extension in (".docx", ".doc")
        )
    
    def extract(self, content: bytes, filename: str) -> ExtractionResult:
        ext = Path(filename).suffix.lower()

        if ext == ".doc":
            # Legacy .doc format - limited support
            return self._extract_legacy_doc(content, filename)

        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx not installed. Run: pip install python-docx")

        # SECURITY: Check decompression ratio before extraction
        # DOCX is a ZIP file - malicious files could have huge decompression ratios
        compressed_size = len(content)
        self._check_decompression_size(compressed_size, filename)

        doc = Document(io.BytesIO(content))

        paragraphs = []
        total_chars = 0
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
                total_chars += len(text)
                # SECURITY: Early termination if extraction is too large
                if total_chars > MAX_DECOMPRESSED_SIZE:
                    raise ValueError(
                        f"Decompression bomb detected: extracted content exceeds "
                        f"{MAX_DECOMPRESSED_SIZE // (1024*1024)}MB limit"
                    )

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                        total_chars += len(cell_text)
                if row_text:
                    paragraphs.append(" | ".join(row_text))
                # SECURITY: Check size during table extraction too
                if total_chars > MAX_DECOMPRESSED_SIZE:
                    raise ValueError(
                        f"Decompression bomb detected: extracted content exceeds "
                        f"{MAX_DECOMPRESSED_SIZE // (1024*1024)}MB limit"
                    )

        return ExtractionResult(
            text="\n\n".join(paragraphs),
            pages=1,  # DOCX doesn't have fixed pages
        )

    def _check_decompression_size(self, compressed_size: int, filename: str) -> None:
        """Check that compressed file isn't suspiciously small (potential zip bomb)."""
        # Very small compressed files that claim to be documents are suspicious
        if compressed_size < 100:
            logger.warning(f"Suspiciously small DOCX file: {filename} ({compressed_size} bytes)")
            # Allow but log - very small files might be legitimate empty docs
    
    def _extract_legacy_doc(self, content: bytes, filename: str) -> ExtractionResult:
        """
        Extract from legacy .doc format.
        
        Limited support - may not work for all files.
        """
        # Try to extract raw text (very basic)
        # For proper .doc support, would need antiword or similar
        try:
            # Attempt basic text extraction
            text = content.decode("latin-1", errors="ignore")
            # Filter to printable characters
            printable = "".join(
                c if c.isprintable() or c in "\n\r\t" else " " 
                for c in text
            )
            # Clean up whitespace
            lines = [line.strip() for line in printable.split("\n")]
            lines = [line for line in lines if line and len(line) > 3]
            
            return ExtractionResult(
                text="\n".join(lines),
                pages=1,
                warnings=["Legacy .doc format - extraction may be incomplete"],
            )
        except Exception as e:
            return ExtractionResult(
                text="",
                pages=1,
                warnings=[f"Failed to extract from legacy .doc: {e}"],
            )


class XLSXExtractor(BaseExtractor):
    """Spreadsheet extractor for XLSX, XLS, and CSV."""
    
    def can_handle(self, content_type: str, extension: str) -> bool:
        return (
            content_type in (
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                "application/vnd.ms-excel",
                "text/csv",
            ) or 
            extension in (".xlsx", ".xls", ".csv", ".tsv")
        )
    
    def extract(self, content: bytes, filename: str) -> ExtractionResult:
        ext = Path(filename).suffix.lower()
        
        if ext == ".csv":
            return self._extract_csv(content, ",")
        elif ext == ".tsv":
            return self._extract_csv(content, "\t")
        elif ext == ".xls":
            return self._extract_xls(content, filename)
        else:
            return self._extract_xlsx(content, filename)
    
    def _extract_csv(self, content: bytes, delimiter: str) -> ExtractionResult:
        """Extract from CSV/TSV."""
        # Try common encodings
        text_content = None
        for encoding in ["utf-8", "utf-8-sig", "latin-1", "cp1252"]:
            try:
                text_content = content.decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        
        if text_content is None:
            return ExtractionResult(
                text="",
                warnings=["Failed to decode CSV file"],
            )
        
        rows = []
        reader = csv.reader(io.StringIO(text_content), delimiter=delimiter)
        for row in reader:
            if any(cell.strip() for cell in row):
                rows.append(" | ".join(cell.strip() for cell in row if cell.strip()))
        
        return ExtractionResult(
            text="\n".join(rows),
            pages=1,
        )
    
    def _extract_xlsx(self, content: bytes, filename: str) -> ExtractionResult:
        """Extract from XLSX."""
        try:
            from openpyxl import load_workbook
        except ImportError:
            raise ImportError("openpyxl not installed. Run: pip install openpyxl")

        # SECURITY: Track total extracted size to prevent decompression bombs
        # XLSX is a ZIP file that could have malicious compression ratios
        compressed_size = len(content)

        wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)

        all_text = []
        warnings = []
        total_chars = 0

        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            sheet_rows = []
            row_count = 0

            for row in sheet.iter_rows(values_only=True):
                # Limit rows per sheet to prevent DoS
                if row_count >= MAX_SPREADSHEET_ROWS:
                    warnings.append(f"Sheet '{sheet_name}' truncated at {MAX_SPREADSHEET_ROWS} rows")
                    break
                row_count += 1

                cells = [str(cell).strip() for cell in row if cell is not None]
                if cells:
                    row_text = " | ".join(cells)
                    sheet_rows.append(row_text)
                    total_chars += len(row_text)

                    # SECURITY: Check for decompression bomb
                    if total_chars > MAX_DECOMPRESSED_SIZE:
                        wb.close()
                        raise ValueError(
                            f"Decompression bomb detected: extracted content exceeds "
                            f"{MAX_DECOMPRESSED_SIZE // (1024*1024)}MB limit"
                        )

            if sheet_rows:
                all_text.append(f"[Sheet: {sheet_name}]")
                all_text.extend(sheet_rows)
                all_text.append("")

        wb.close()

        # SECURITY: Final check on extraction ratio
        if compressed_size > 0 and total_chars > 0:
            ratio = total_chars / compressed_size
            if ratio > MAX_EXTRACTION_RATIO:
                logger.warning(
                    f"High extraction ratio for {filename}: {ratio:.1f}x "
                    f"({compressed_size} bytes -> {total_chars} chars)"
                )

        return ExtractionResult(
            text="\n".join(all_text),
            pages=len(wb.sheetnames),
            warnings=warnings,
        )

    def _extract_xls(self, content: bytes, filename: str) -> ExtractionResult:
        """Extract from legacy XLS."""
        try:
            import xlrd
        except ImportError:
            raise ImportError("xlrd not installed. Run: pip install xlrd")
        
        wb = xlrd.open_workbook(file_contents=content)

        all_text = []
        warnings = []
        for sheet_idx in range(wb.nsheets):
            sheet = wb.sheet_by_index(sheet_idx)
            sheet_rows = []

            # Limit rows per sheet to prevent DoS
            max_rows = min(sheet.nrows, MAX_SPREADSHEET_ROWS)
            if sheet.nrows > MAX_SPREADSHEET_ROWS:
                warnings.append(f"Sheet '{sheet.name}' truncated at {MAX_SPREADSHEET_ROWS} rows")

            for row_idx in range(max_rows):
                row = sheet.row_values(row_idx)
                cells = [str(cell).strip() for cell in row if cell]
                if cells:
                    sheet_rows.append(" | ".join(cells))

            if sheet_rows:
                all_text.append(f"[Sheet: {sheet.name}]")
                all_text.extend(sheet_rows)
                all_text.append("")

        return ExtractionResult(
            text="\n".join(all_text),
            pages=wb.nsheets,
            warnings=warnings,
        )


class ImageExtractor(BaseExtractor):
    """
    Image text extractor using OCR with document intelligence.
    
    Handles JPEG, PNG, TIFF, HEIC, GIF, BMP, WebP.
    For multi-page TIFFs, saves individual pages to temp files.
    
    Uses EnhancedOCRProcessor for:
    - Document type detection (driver's license, insurance card, etc.)
    - Layout-aware text spacing
    - Pre-extraction of PHI fields with HIPAA categories
    """
    
    def __init__(
        self, 
        ocr_engine: Optional["OCREngine"] = None,
        temp_dir: Optional["SecureTempDir"] = None,
        enable_enhanced_processing: bool = True,
    ):
        """
        Initialize image extractor.
        
        Args:
            ocr_engine: OCR engine for text extraction
            temp_dir: Secure temp directory for page images (optional)
            enable_enhanced_processing: Use EnhancedOCRProcessor for document intelligence
        """
        self.ocr_engine = ocr_engine
        self.temp_dir = temp_dir
        self.enable_enhanced_processing = enable_enhanced_processing
        self._enhanced_processor = None
    
    @property
    def enhanced_processor(self):
        """Lazy-load EnhancedOCRProcessor."""
        if self._enhanced_processor is None and self.enable_enhanced_processing:
            try:
                from .enhanced_ocr import EnhancedOCRProcessor
                self._enhanced_processor = EnhancedOCRProcessor()
                logger.info("EnhancedOCRProcessor initialized for document intelligence")
            except Exception as e:
                logger.warning(f"Could not initialize EnhancedOCRProcessor: {e}")
                self.enable_enhanced_processing = False
        return self._enhanced_processor
    
    def can_handle(self, content_type: str, extension: str) -> bool:
        return (
            content_type.startswith("image/") or
            extension in (".jpg", ".jpeg", ".png", ".tiff", ".tif", 
                         ".heic", ".heif", ".gif", ".bmp", ".webp")
        )
    
    def extract(
        self, 
        content: bytes, 
        filename: str,
        save_pages: bool = True,
    ) -> ExtractionResult:
        """
        Extract text from image with document intelligence.
        
        Args:
            content: Image bytes
            filename: Original filename
            save_pages: If True, save page images to temp directory
                       
        Returns:
            ExtractionResult with OCR data, document type, and pre-extracted PHI
        """
        if not self.ocr_engine or not self.ocr_engine.is_available:
            return ExtractionResult(
                text="",
                pages=1,
                needs_ocr=True,
                warnings=["OCR engine not available for image extraction"],
            )
        
        ext = Path(filename).suffix.lower()
        
        try:
            from PIL import Image
            import numpy as np
            
            # Handle HEIC format
            if ext in (".heic", ".heif"):
                try:
                    from pillow_heif import register_heif_opener
                    register_heif_opener()
                except ImportError:
                    return ExtractionResult(
                        text="",
                        pages=1,
                        warnings=["HEIC support not available. Run: pip install pillow-heif"],
                    )
            
            # Handle multi-page TIFF
            if ext in (".tiff", ".tif"):
                return self._extract_multipage_tiff(content, filename, save_pages)
            
            # Single image
            img = Image.open(io.BytesIO(content))
            
            # Convert to RGB if necessary
            if img.mode not in ("RGB", "L"):
                img = img.convert("RGB")
            
            # Convert to numpy array for enhanced processing
            img_array = np.array(img)
            
            # Save to temp if requested
            temp_path = None
            if save_pages and self.temp_dir:
                img_buffer = io.BytesIO()
                img.save(img_buffer, format='PNG')
                temp_path_obj = self.temp_dir.write_page(0, img_buffer.getvalue())
                temp_path = str(temp_path_obj)
            
            # Get raw OCR with coordinates
            ocr_result = self.ocr_engine.extract_with_coordinates(img)
            
            # Apply enhanced processing for document intelligence
            document_type = None
            is_id_document = False
            phi_fields = None
            enhanced_text = None
            enhancements = []
            
            if self.enable_enhanced_processing and self.enhanced_processor and ocr_result.blocks:
                try:
                    enhanced_result = self.enhanced_processor.process(
                        image=img_array,
                        ocr_result=ocr_result,
                        apply_document_cleaning=True,
                    )
                    
                    document_type = enhanced_result.document_type.name
                    is_id_document = enhanced_result.is_id_card
                    phi_fields = enhanced_result.phi_fields
                    enhanced_text = enhanced_result.enhanced_text
                    enhancements = enhanced_result.enhancements_applied
                    
                    logger.info(
                        f"Document intelligence: type={document_type}, "
                        f"is_id={is_id_document}, "
                        f"phi_fields={len(phi_fields) if phi_fields else 0}, "
                        f"enhancements={enhancements}"
                    )
                    
                except Exception as e:
                    logger.warning(f"Enhanced OCR processing failed, using raw OCR: {e}")
                    enhancements.append(f"enhanced_failed:{str(e)[:50]}")
            
            page_info = PageInfo(
                page_num=0,
                text=enhanced_text or ocr_result.full_text,
                is_scanned=True,
                ocr_result=ocr_result,
                temp_image_path=temp_path,
            )
            
            return ExtractionResult(
                text=ocr_result.full_text,  # Raw text for offset mapping
                pages=1,
                needs_ocr=True,
                ocr_pages=[0],
                confidence=ocr_result.confidence,
                ocr_results=[ocr_result],
                page_infos=[page_info],
                temp_dir_path=str(self.temp_dir.path) if self.temp_dir and self.temp_dir.path else None,
                # Document intelligence
                document_type=document_type,
                is_id_document=is_id_document,
                phi_fields=phi_fields,
                enhanced_text=enhanced_text,
                enhancements_applied=enhancements,
            )
            
        except Exception as e:
            logger.error(f"Image extraction failed: {e}")
            return ExtractionResult(
                text="",
                pages=1,
                warnings=[f"Image extraction failed: {e}"],
            )
    
    def _extract_multipage_tiff(
        self, 
        content: bytes, 
        filename: str,
        save_pages: bool = True,
    ) -> ExtractionResult:
        """Extract text from multi-page TIFF with document intelligence."""
        from PIL import Image
        import numpy as np
        
        img = Image.open(io.BytesIO(content))
        
        pages_text = []
        page_infos = []
        ocr_results = []
        confidences = []
        
        # Document intelligence (from first page)
        document_type = None
        is_id_document = False
        phi_fields = None
        enhanced_texts = []
        enhancements = []
        
        try:
            page_num = 0
            while True:
                # Early exit if page limit exceeded (prevents DoS via large TIFFs)
                if page_num >= MAX_DOCUMENT_PAGES:
                    logger.warning(f"TIFF exceeds {MAX_DOCUMENT_PAGES} page limit, truncating")
                    break

                img.seek(page_num)

                # Convert to RGB
                frame = img.convert("RGB")
                frame_array = np.array(frame)
                
                # Save to temp if requested
                temp_path = None
                if save_pages and self.temp_dir:
                    img_buffer = io.BytesIO()
                    frame.save(img_buffer, format='PNG')
                    temp_path_obj = self.temp_dir.write_page(page_num, img_buffer.getvalue())
                    temp_path = str(temp_path_obj)
                
                # Get raw OCR with coordinates
                ocr_result = self.ocr_engine.extract_with_coordinates(frame)
                
                # Apply enhanced processing (document detection on first page)
                enhanced_text = ocr_result.full_text
                
                if self.enable_enhanced_processing and self.enhanced_processor and ocr_result.blocks:
                    try:
                        enhanced_result = self.enhanced_processor.process(
                            image=frame_array,
                            ocr_result=ocr_result,
                            apply_document_cleaning=True,
                        )
                        
                        enhanced_text = enhanced_result.enhanced_text
                        
                        # Capture document type from first page
                        if page_num == 0:
                            document_type = enhanced_result.document_type.name
                            is_id_document = enhanced_result.is_id_card
                            phi_fields = enhanced_result.phi_fields
                            enhancements = enhanced_result.enhancements_applied
                        
                    except Exception as e:
                        logger.warning(f"Enhanced processing failed for page {page_num}: {e}")
                
                pages_text.append(ocr_result.full_text)
                enhanced_texts.append(enhanced_text)
                ocr_results.append(ocr_result)
                confidences.append(ocr_result.confidence)
                
                page_infos.append(PageInfo(
                    page_num=page_num,
                    text=enhanced_text,
                    is_scanned=True,
                    ocr_result=ocr_result,
                    temp_image_path=temp_path,
                ))
                
                page_num += 1
                
        except EOFError:
            # End of pages
            pass
        
        avg_confidence = sum(confidences) / len(confidences) if confidences else 0.0
        
        return ExtractionResult(
            text="\n\n".join(pages_text),
            pages=len(pages_text),
            needs_ocr=True,
            ocr_pages=list(range(len(pages_text))),
            confidence=avg_confidence,
            ocr_results=ocr_results,
            page_infos=page_infos,
            temp_dir_path=str(self.temp_dir.path) if self.temp_dir and self.temp_dir.path else None,
            # Document intelligence (from first page)
            document_type=document_type,
            is_id_document=is_id_document,
            phi_fields=phi_fields,
            enhanced_text="\n\n".join(enhanced_texts) if enhanced_texts else None,
            enhancements_applied=enhancements,
        )


class TextExtractor(BaseExtractor):
    """Plain text file extractor."""
    
    def can_handle(self, content_type: str, extension: str) -> bool:
        return (
            content_type == "text/plain" or
            extension == ".txt"
        )
    
    def extract(self, content: bytes, filename: str) -> ExtractionResult:
        # Try common encodings
        for encoding in ["utf-8", "utf-8-sig", "latin-1", "cp1252"]:
            try:
                text = content.decode(encoding)
                return ExtractionResult(
                    text=text,
                    pages=1,
                )
            except UnicodeDecodeError:
                continue
        
        return ExtractionResult(
            text="",
            pages=1,
            warnings=["Failed to decode text file"],
        )


class RTFExtractor(BaseExtractor):
    """RTF document extractor using striprtf."""
    
    def can_handle(self, content_type: str, extension: str) -> bool:
        return (
            content_type == "application/rtf" or
            extension == ".rtf"
        )
    
    def extract(self, content: bytes, filename: str) -> ExtractionResult:
        try:
            from striprtf.striprtf import rtf_to_text
        except ImportError:
            raise ImportError("striprtf not installed. Run: pip install striprtf")
        
        # Try to decode
        for encoding in ["utf-8", "latin-1", "cp1252"]:
            try:
                rtf_content = content.decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        else:
            return ExtractionResult(
                text="",
                pages=1,
                warnings=["Failed to decode RTF file"],
            )
        
        try:
            text = rtf_to_text(rtf_content)
            return ExtractionResult(
                text=text,
                pages=1,
            )
        except Exception as e:
            return ExtractionResult(
                text="",
                pages=1,
                warnings=[f"RTF extraction failed: {e}"],
            )

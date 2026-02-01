"""Office document extractors (Word, Excel, RTF, Text)."""

import csv
import io
import logging
from pathlib import Path

from ..constants import (
    MAX_DECOMPRESSED_SIZE,
    MAX_EXTRACTION_RATIO,
    MAX_SPREADSHEET_ROWS,
)
from .base import BaseExtractor, ExtractionResult

logger = logging.getLogger(__name__)


class DOCXExtractor(BaseExtractor):
    """Word document extractor using python-docx."""

    def can_handle(self, content_type: str, extension: str) -> bool:
        return (
            content_type in (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "application/msword",
            ) or
            extension in (".docx", ".doc")
        )

    def extract(self, content: bytes, filename: str) -> ExtractionResult:
        ext = Path(filename).suffix.lower()

        if ext == ".doc":
            return self._extract_legacy_doc(content, filename)

        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx not installed. Run: pip install python-docx")

        compressed_size = len(content)
        self._check_decompression_size(compressed_size, filename)

        doc = Document(io.BytesIO(content))

        paragraphs = []
        total_chars = 0
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
                total_chars += len(text)
                if total_chars > MAX_DECOMPRESSED_SIZE:
                    raise ValueError(
                        f"Decompression bomb detected: extracted content exceeds "
                        f"{MAX_DECOMPRESSED_SIZE // (1024*1024)}MB limit"
                    )

        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                        total_chars += len(cell_text)
                if row_text:
                    paragraphs.append(" | ".join(row_text))
                if total_chars > MAX_DECOMPRESSED_SIZE:
                    raise ValueError(
                        f"Decompression bomb detected: extracted content exceeds "
                        f"{MAX_DECOMPRESSED_SIZE // (1024*1024)}MB limit"
                    )

        return ExtractionResult(
            text="\n\n".join(paragraphs),
            pages=1,
        )

    def _check_decompression_size(self, compressed_size: int, filename: str) -> None:
        if compressed_size < 100:
            logger.warning(f"Suspiciously small DOCX file: {filename} ({compressed_size} bytes)")

    def _extract_legacy_doc(self, content: bytes, filename: str) -> ExtractionResult:
        try:
            text = content.decode("latin-1", errors="ignore")
            printable = "".join(
                c if c.isprintable() or c in "\n\r\t" else " "
                for c in text
            )
            lines = [line.strip() for line in printable.split("\n")]
            lines = [line for line in lines if line and len(line) > 3]

            return ExtractionResult(
                text="\n".join(lines),
                pages=1,
                warnings=["Legacy .doc format - extraction may be incomplete"],
            )
        except (ValueError, UnicodeDecodeError) as e:
            return ExtractionResult(
                text="",
                pages=1,
                warnings=[f"Failed to extract from legacy .doc: {e}"],
            )


class XLSXExtractor(BaseExtractor):
    """Spreadsheet extractor for XLSX, XLS, and CSV."""

    def can_handle(self, content_type: str, extension: str) -> bool:
        return (
            content_type in (
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                "application/vnd.ms-excel",
                "text/csv",
            ) or
            extension in (".xlsx", ".xls", ".csv", ".tsv")
        )

    def extract(self, content: bytes, filename: str) -> ExtractionResult:
        ext = Path(filename).suffix.lower()

        if ext == ".csv":
            return self._extract_csv(content, ",")
        elif ext == ".tsv":
            return self._extract_csv(content, "\t")
        elif ext == ".xls":
            return self._extract_xls(content, filename)
        else:
            return self._extract_xlsx(content, filename)

    def _extract_csv(self, content: bytes, delimiter: str) -> ExtractionResult:
        text_content = None
        for encoding in ["utf-8", "utf-8-sig", "latin-1", "cp1252"]:
            try:
                text_content = content.decode(encoding)
                break
            except UnicodeDecodeError:
                continue

        if text_content is None:
            return ExtractionResult(
                text="",
                warnings=["Failed to decode CSV file"],
            )

        rows = []
        warnings = []
        row_count = 0
        reader = csv.reader(io.StringIO(text_content), delimiter=delimiter)
        for row in reader:
            if row_count >= MAX_SPREADSHEET_ROWS:  # HIGH-003: prevent OOM
                warnings.append(f"CSV truncated at {MAX_SPREADSHEET_ROWS} rows")
                break
            row_count += 1
            if any(cell.strip() for cell in row):
                rows.append(" | ".join(cell.strip() for cell in row if cell.strip()))

        return ExtractionResult(
            text="\n".join(rows),
            pages=1,
            warnings=warnings,
        )

    def _extract_xlsx(self, content: bytes, filename: str) -> ExtractionResult:
        try:
            from openpyxl import load_workbook
        except ImportError:
            raise ImportError("openpyxl not installed. Run: pip install openpyxl")

        compressed_size = len(content)
        wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)

        all_text = []
        warnings = []
        total_chars = 0

        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            sheet_rows = []
            row_count = 0

            for row in sheet.iter_rows(values_only=True):
                if row_count >= MAX_SPREADSHEET_ROWS:
                    warnings.append(f"Sheet '{sheet_name}' truncated at {MAX_SPREADSHEET_ROWS} rows")
                    break
                row_count += 1

                cells = [str(cell).strip() for cell in row if cell is not None]
                if cells:
                    row_text = " | ".join(cells)
                    sheet_rows.append(row_text)
                    total_chars += len(row_text)

                    if total_chars > MAX_DECOMPRESSED_SIZE:
                        wb.close()
                        raise ValueError(
                            f"Decompression bomb detected: extracted content exceeds "
                            f"{MAX_DECOMPRESSED_SIZE // (1024*1024)}MB limit"
                        )

            if sheet_rows:
                all_text.append(f"[Sheet: {sheet_name}]")
                all_text.extend(sheet_rows)
                all_text.append("")

        wb.close()

        if compressed_size > 0 and total_chars > 0:
            ratio = total_chars / compressed_size
            if ratio > MAX_EXTRACTION_RATIO:  # HIGH-005: decompression bomb
                raise ValueError(
                    f"Suspected decompression bomb: extraction ratio {ratio:.1f}x exceeds "
                    f"maximum {MAX_EXTRACTION_RATIO}x for {filename} "
                    f"({compressed_size} bytes -> {total_chars} chars)"
                )

        return ExtractionResult(
            text="\n".join(all_text),
            pages=len(wb.sheetnames),
            warnings=warnings,
        )

    def _extract_xls(self, content: bytes, filename: str) -> ExtractionResult:
        try:
            import xlrd
        except ImportError:
            raise ImportError("xlrd not installed. Run: pip install xlrd")

        wb = xlrd.open_workbook(file_contents=content)

        all_text = []
        warnings = []
        for sheet_idx in range(wb.nsheets):
            sheet = wb.sheet_by_index(sheet_idx)
            sheet_rows = []

            max_rows = min(sheet.nrows, MAX_SPREADSHEET_ROWS)
            if sheet.nrows > MAX_SPREADSHEET_ROWS:
                warnings.append(f"Sheet '{sheet.name}' truncated at {MAX_SPREADSHEET_ROWS} rows")

            for row_idx in range(max_rows):
                row = sheet.row_values(row_idx)
                cells = [str(cell).strip() for cell in row if cell]
                if cells:
                    sheet_rows.append(" | ".join(cells))

            if sheet_rows:
                all_text.append(f"[Sheet: {sheet.name}]")
                all_text.extend(sheet_rows)
                all_text.append("")

        return ExtractionResult(
            text="\n".join(all_text),
            pages=wb.nsheets,
            warnings=warnings,
        )


class TextExtractor(BaseExtractor):
    """Plain text file extractor."""

    def can_handle(self, content_type: str, extension: str) -> bool:
        return content_type == "text/plain" or extension == ".txt"

    def extract(self, content: bytes, filename: str) -> ExtractionResult:
        for encoding in ["utf-8", "utf-8-sig", "latin-1", "cp1252"]:
            try:
                text = content.decode(encoding)
                return ExtractionResult(
                    text=text,
                    pages=1,
                )
            except UnicodeDecodeError:
                continue

        return ExtractionResult(
            text="",
            pages=1,
            warnings=["Failed to decode text file"],
        )


class RTFExtractor(BaseExtractor):
    """RTF document extractor using striprtf."""

    def can_handle(self, content_type: str, extension: str) -> bool:
        return content_type == "application/rtf" or extension == ".rtf"

    def extract(self, content: bytes, filename: str) -> ExtractionResult:
        try:
            from striprtf.striprtf import rtf_to_text
        except ImportError:
            raise ImportError("striprtf not installed. Run: pip install striprtf")

        for encoding in ["utf-8", "latin-1", "cp1252"]:
            try:
                rtf_content = content.decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        else:
            return ExtractionResult(
                text="",
                pages=1,
                warnings=["Failed to decode RTF file"],
            )

        try:
            text = rtf_to_text(rtf_content)
            return ExtractionResult(
                text=text,
                pages=1,
            )
        except (ValueError, RuntimeError) as e:
            return ExtractionResult(
                text="",
                pages=1,
                warnings=[f"RTF extraction failed: {e}"],
            )

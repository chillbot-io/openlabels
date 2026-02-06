"""Tests for file extractors.

Adapted from openrisk tests for openlabels extractors module.
"""

import io
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest


# =============================================================================
# Dependency Check Helpers
# =============================================================================

def has_pymupdf():
    """Check if PyMuPDF is available."""
    try:
        import fitz
        return True
    except ImportError:
        return False


def has_docx():
    """Check if python-docx is available."""
    try:
        import docx
        return True
    except ImportError:
        return False


def has_openpyxl():
    """Check if openpyxl is available."""
    try:
        import openpyxl
        return True
    except ImportError:
        return False


def has_pillow():
    """Check if Pillow is available."""
    try:
        from PIL import Image
        return True
    except ImportError:
        return False


# Skip markers
requires_pymupdf = pytest.mark.skipif(not has_pymupdf(), reason="PyMuPDF not installed")
requires_docx = pytest.mark.skipif(not has_docx(), reason="python-docx not installed")
requires_openpyxl = pytest.mark.skipif(not has_openpyxl(), reason="openpyxl not installed")
requires_pillow = pytest.mark.skipif(not has_pillow(), reason="Pillow not installed")


# =============================================================================
# ExtractionResult Tests
# =============================================================================

class TestExtractionResult:
    """Tests for ExtractionResult dataclass."""

    def test_result_with_warnings(self):
        """Test result with warnings."""
        from openlabels.core.extractors import ExtractionResult

        result = ExtractionResult(
            text="Partial text",
            pages=10,
            warnings=["Document truncated at 10 pages"],
        )

        assert len(result.warnings) == 1
        assert "truncated" in result.warnings[0]

    def test_result_with_ocr_info(self):
        """Test result with OCR information."""
        from openlabels.core.extractors import ExtractionResult

        result = ExtractionResult(
            text="OCR extracted text",
            pages=2,
            needs_ocr=True,
            ocr_pages=[0, 1],
            confidence=0.85,
        )

        assert result.needs_ocr is True
        assert result.ocr_pages == [0, 1]
        assert result.confidence == 0.85



# =============================================================================
# PDF Extractor Tests
# =============================================================================

@requires_pymupdf
class TestPDFExtractor:
    """Tests for PDFExtractor."""

    def test_can_handle_pdf_content_type(self):
        """Test that PDFExtractor handles PDF content type."""
        from openlabels.core.extractors import PDFExtractor

        extractor = PDFExtractor()
        assert extractor.can_handle("application/pdf", ".txt") is True

    def test_can_handle_pdf_extension(self):
        """Test that PDFExtractor handles .pdf extension."""
        from openlabels.core.extractors import PDFExtractor

        extractor = PDFExtractor()
        assert extractor.can_handle("text/plain", ".pdf") is True

    def test_does_not_handle_other_types(self):
        """Test that PDFExtractor rejects non-PDF types."""
        from openlabels.core.extractors import PDFExtractor

        extractor = PDFExtractor()
        assert extractor.can_handle("text/plain", ".txt") is False
        assert extractor.can_handle("image/png", ".png") is False

    def test_extract_simple_pdf(self):
        """Test extracting text from a simple PDF."""
        import fitz
        from openlabels.core.extractors import PDFExtractor

        # Create a simple PDF in memory
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 50), "Test content with SSN: 123-45-6789")
        pdf_bytes = doc.tobytes()
        doc.close()

        extractor = PDFExtractor()
        result = extractor.extract(pdf_bytes, "test.pdf")

        assert result.text is not None
        # PDF extraction should capture the text we inserted
        assert "Test content" in result.text, \
            f"Expected 'Test content' in extracted text, got: {result.text[:200]}"

    def test_extract_multi_page_pdf(self):
        """Test extracting from multi-page PDF."""
        import fitz
        from openlabels.core.extractors import PDFExtractor

        doc = fitz.open()
        for i in range(3):
            page = doc.new_page()
            page.insert_text((50, 50), f"Page {i + 1} content")
        pdf_bytes = doc.tobytes()
        doc.close()

        extractor = PDFExtractor()
        result = extractor.extract(pdf_bytes, "multi.pdf")

        assert result.text is not None
        assert result.pages == 3

    def test_invalid_pdf_raises_error(self):
        """Test that invalid PDF raises appropriate error."""
        from openlabels.core.extractors import PDFExtractor

        extractor = PDFExtractor()

        with pytest.raises(Exception):
            extractor.extract(b"This is not a PDF", "fake.pdf")


# =============================================================================
# DOCX Extractor Tests
# =============================================================================

@requires_docx
class TestDOCXExtractor:
    """Tests for DOCXExtractor."""

    def test_can_handle_docx(self):
        """Test that DOCXExtractor handles DOCX."""
        from openlabels.core.extractors import DOCXExtractor

        extractor = DOCXExtractor()
        assert extractor.can_handle(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".docx"
        ) is True

    def test_does_not_handle_pdf(self):
        """Test that DOCXExtractor rejects PDF."""
        from openlabels.core.extractors import DOCXExtractor

        extractor = DOCXExtractor()
        assert extractor.can_handle("application/pdf", ".pdf") is False

    def test_extract_simple_docx(self):
        """Test extracting text from a simple DOCX."""
        from docx import Document
        from openlabels.core.extractors import DOCXExtractor

        doc = Document()
        doc.add_paragraph("Test paragraph with email: test@example.com")
        doc.add_paragraph("Second paragraph with phone: 555-123-4567")

        buffer = io.BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()

        extractor = DOCXExtractor()
        result = extractor.extract(docx_bytes, "test.docx")

        # DOCX extraction should capture the paragraph content
        assert "Test paragraph" in result.text, \
            f"Expected 'Test paragraph' in extracted text, got: {result.text[:200]}"

    def test_extract_docx_with_tables(self):
        """Test extracting text from DOCX with tables."""
        from docx import Document
        from openlabels.core.extractors import DOCXExtractor

        doc = Document()
        doc.add_paragraph("Header")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Name"
        table.cell(0, 1).text = "SSN"
        table.cell(1, 0).text = "John Doe"
        table.cell(1, 1).text = "123-45-6789"

        buffer = io.BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()

        extractor = DOCXExtractor()
        result = extractor.extract(docx_bytes, "table.docx")

        # DOCX extraction should capture table cell content
        assert "John Doe" in result.text, \
            f"Expected 'John Doe' from table in extracted text, got: {result.text[:200]}"
        assert "123-45-6789" in result.text, \
            f"Expected SSN from table in extracted text, got: {result.text[:200]}"


# =============================================================================
# XLSX Extractor Tests
# =============================================================================

@requires_openpyxl
class TestXLSXExtractor:
    """Tests for XLSXExtractor."""

    def test_can_handle_xlsx(self):
        """Test that XLSXExtractor handles XLSX."""
        from openlabels.core.extractors import XLSXExtractor

        extractor = XLSXExtractor()
        assert extractor.can_handle(
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            ".xlsx"
        ) is True

    def test_can_handle_csv(self):
        """Test that XLSXExtractor handles CSV."""
        from openlabels.core.extractors import XLSXExtractor

        extractor = XLSXExtractor()
        assert extractor.can_handle("text/csv", ".csv") is True

    def test_extract_simple_xlsx(self):
        """Test extracting text from a simple XLSX."""
        from openpyxl import Workbook
        from openlabels.core.extractors import XLSXExtractor

        wb = Workbook()
        ws = wb.active
        ws['A1'] = "Name"
        ws['B1'] = "SSN"
        ws['A2'] = "John Doe"
        ws['B2'] = "123-45-6789"

        buffer = io.BytesIO()
        wb.save(buffer)
        xlsx_bytes = buffer.getvalue()

        extractor = XLSXExtractor()
        result = extractor.extract(xlsx_bytes, "test.xlsx")

        assert result is not None
        # XLSX extraction should capture cell content
        assert "John Doe" in result.text, \
            f"Expected 'John Doe' from cell in extracted text, got: {result.text[:200]}"

    def test_extract_csv(self):
        """Test extracting text from CSV."""
        from openlabels.core.extractors import XLSXExtractor

        csv_content = b"Name,Email\nJohn,john@test.com\nJane,jane@test.com"

        extractor = XLSXExtractor()
        result = extractor.extract(csv_content, "test.csv")

        # CSV extraction should capture the content
        assert "john@test.com" in result.text, \
            f"Expected email in extracted CSV text, got: {result.text[:200]}"


# =============================================================================
# Image Extractor Tests
# =============================================================================

@requires_pillow
class TestImageExtractor:
    """Tests for ImageExtractor."""

    def test_can_handle_png(self):
        """Test that ImageExtractor handles PNG."""
        from openlabels.core.extractors import ImageExtractor

        extractor = ImageExtractor()
        assert extractor.can_handle("image/png", ".png") is True

    def test_can_handle_jpeg(self):
        """Test that ImageExtractor handles JPEG."""
        from openlabels.core.extractors import ImageExtractor

        extractor = ImageExtractor()
        assert extractor.can_handle("image/jpeg", ".jpg") is True
        assert extractor.can_handle("image/jpeg", ".jpeg") is True

    def test_does_not_handle_pdf(self):
        """Test that ImageExtractor rejects PDF."""
        from openlabels.core.extractors import ImageExtractor

        extractor = ImageExtractor()
        assert extractor.can_handle("application/pdf", ".pdf") is False

    def test_extract_without_ocr(self):
        """Test extracting from image without OCR returns empty."""
        from PIL import Image
        from openlabels.core.extractors import ImageExtractor

        img = Image.new('RGB', (100, 100), color='white')
        buffer = io.BytesIO()
        img.save(buffer, format='PNG')
        img_bytes = buffer.getvalue()

        extractor = ImageExtractor()  # No OCR engine
        result = extractor.extract(img_bytes, "test.png")

        # Without OCR, should return empty or warning
        assert result is not None


# =============================================================================
# Text Extractor Tests
# =============================================================================

class TestTextExtractor:
    """Tests for TextExtractor."""

    def test_can_handle_txt(self):
        """Test that TextExtractor handles .txt."""
        from openlabels.core.extractors import TextExtractor

        extractor = TextExtractor()
        assert extractor.can_handle("text/plain", ".txt") is True

    def test_extract_utf8(self):
        """Test extracting UTF-8 text."""
        from openlabels.core.extractors import TextExtractor

        text_content = "Hello, this is a test with special chars: éàü"
        text_bytes = text_content.encode("utf-8")

        extractor = TextExtractor()
        result = extractor.extract(text_bytes, "test.txt")

        assert result.text == text_content

    def test_extract_latin1(self):
        """Test extracting Latin-1 text."""
        from openlabels.core.extractors import TextExtractor

        text_content = "Hello, this is a test"
        text_bytes = text_content.encode("latin-1")

        extractor = TextExtractor()
        result = extractor.extract(text_bytes, "test.txt")

        assert "Hello" in result.text


# =============================================================================
# RTF Extractor Tests
# =============================================================================

class TestRTFExtractor:
    """Tests for RTFExtractor."""

    def test_can_handle_rtf(self):
        """Test that RTFExtractor handles RTF."""
        from openlabels.core.extractors import RTFExtractor

        extractor = RTFExtractor()
        assert extractor.can_handle("application/rtf", ".rtf") is True

    def test_does_not_handle_docx(self):
        """Test that RTFExtractor rejects DOCX."""
        from openlabels.core.extractors import RTFExtractor

        extractor = RTFExtractor()
        assert extractor.can_handle("application/docx", ".docx") is False


# =============================================================================
# Registry Tests
# =============================================================================

class TestExtractorRegistry:
    """Tests for extractor registry functions."""

    def test_get_extractor_pdf(self):
        """Test getting PDF extractor."""
        from openlabels.core.extractors import get_extractor

        extractor = get_extractor("application/pdf", ".pdf")
        assert extractor is not None

    def test_get_extractor_unknown(self):
        """Test getting extractor for unknown type."""
        from openlabels.core.extractors import get_extractor

        extractor = get_extractor("application/x-unknown", ".xyz")
        assert extractor is None

    def test_extract_text_function(self):
        """Test extract_text convenience function."""
        from openlabels.core.extractors import extract_text

        # Plain text should work
        result = extract_text(b"Hello world", "test.txt")
        assert "Hello" in result.text

    def test_extract_text_unknown_type(self):
        """Test extract_text with unknown file type."""
        from openlabels.core.extractors import extract_text

        result = extract_text(b"some content", "test.xyz")
        # Should return result with warning
        assert result is not None
        assert len(result.warnings) > 0


# =============================================================================
# Security Tests
# =============================================================================

